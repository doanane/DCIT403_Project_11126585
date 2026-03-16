from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3C, 0x6B)
TEAL   = RGBColor(0x0D, 0x6E, 0x6E)
DARK   = RGBColor(0x1C, 0x1C, 0x1C)
GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LBLUE  = RGBColor(0xE8, 0xF0, 0xFE)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val', 'none'))
        tag.set(qn('w:sz'),    kwargs.get('sz',  '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Heading helpers ───────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    p.paragraph_format.keep_with_next = True
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(15)
    run.font.color.rgb = TEAL
    run.font.name = 'Calibri'
    p.paragraph_format.keep_with_next = True
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    p.paragraph_format.keep_with_next = True
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(indent * 0.25)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK
    run.font.name  = 'Calibri'
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK
    run.font.name  = 'Calibri'
    return p

def labelled(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label + ': ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = TEAL
    r1.font.name = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARK
    r2.font.name = 'Calibri'
    return p

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A3C6B')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    return p

def add_table(headers, rows, header_color='1A3C6B'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size  = Pt(10)
        run.font.color.rgb = WHITE
        run.font.name  = 'Calibri'
        p.paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        bg = 'EBF5FB' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = tbl.rows[ri + 1].cells[ci]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('MedStock')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = NAVY
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Hospital Pharmacy Stock Depletion and Emergency Resupply Coordination')
r.font.size = Pt(16)
r.font.color.rgb = TEAL
r.font.name = 'Calibri'
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
r = p.add_run('Multi-Agent System Design and Implementation Report')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = DARK
r.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DCIT 403 — Intelligent Agent Systems')
r.font.size = Pt(12)
r.font.color.rgb = GREY
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prometheus Methodology (Padgham & Winikoff, 2004)')
r.font.size = Pt(12)
r.font.color.rgb = GREY
r.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Student ID: 11126586')
r.font.size = Pt(12)
r.font.color.rgb = GREY
r.font.name = 'Calibri'

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1('1. System Overview')
divider()

body(
    'MedStock is a multi-agent intelligent system designed to solve a persistent operational '
    'failure in hospital pharmacies: critical drug stock depletion goes undetected until a nurse '
    'or doctor physically discovers an empty shelf during an emergency. MedStock deploys five '
    'coordinated agents that continuously monitor drug stock levels across five hospital wards, '
    'classify shortages by clinical severity and drug category, coordinate internal ward '
    'transfers, initiate and track external procurement requests, escalate overdue orders, and '
    'monitor batch expiry with FEFO (First Expired, First Out) enforcement.'
)
body(
    'The system is implemented entirely in Python 3 using the standard library, with a Tkinter '
    'graphical interface for simulation control and visualisation. It follows the Prometheus '
    'methodology through five design phases from system specification to implementation.'
)

h2('Why an Agent-Based Approach?')
bullet('Reactivity: Stock levels change continuously as patients consume medication. Asynchronous sensor readings require immediate classification and response — impossible with a scheduled batch job.', 0)
bullet('Proactivity: The ProcurementEscalationAgent checks all pending orders every cycle without external prompting, escalating overdue procurements automatically.', 0)
bullet('Social ability: Four functional agents mirror real hospital departments (surveillance, assessment, logistics, procurement) and communicate through structured messages.', 0)
bullet('Situatedness: Agents perceive stock readings, act on ward databases, and trigger supplier-facing procurement — embedded in the hospital operational environment.', 0)

h2('Hospital Environment')
add_table(
    ['Item', 'Details'],
    [
        ['Wards monitored',  'ICU, Emergency, Surgical, Maternity, General'],
        ['Drugs tracked',    'Insulin, Morphine, Amoxicillin, Paracetamol, Metformin'],
        ['Suppliers',        'MedPharm Ghana, PharmaControl Ltd, BasicMeds Supplies'],
        ['Simulation steps', '20 discrete steps'],
        ['Entry point',      'main.py → Tkinter GUI'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — PROMETHEUS METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Prometheus Methodology — Five Phases')
divider()
body(
    'The Prometheus methodology (Padgham & Winikoff, 2004) structures agent system design '
    'into four phases: System Specification, Architectural Design, Interaction Design, and '
    'Detailed Design. MedStock adds a fifth phase: the Implementation Report.'
)

# ── Phase 1 ───────────────────────────────────────────────────────────────────
h2('Phase 1 — System Specification')
body(
    'Phase 1 defines the problem, goals, functionalities, scenarios, and environment. '
    'It establishes what the system must achieve before any design decisions are made.'
)
h3('Top-Level Goal')
body(
    'Ensure continuous availability of critical pharmaceutical supplies across all hospital '
    'wards by automatically detecting shortages, coordinating transfers, initiating '
    'procurement, and escalating overdue orders.'
)
h3('Sub-Goals')
add_table(
    ['Goal', 'Description'],
    [
        ['G1 — Monitor stock',       'Receive sensor readings, compute severity percentages, classify as CRITICAL / HIGH / MEDIUM, suppress duplicates'],
        ['G2 — Classify shortages',  'Map shortage to drug category (CONTROLLED / ESSENTIAL / STANDARD), select response plan, prevent duplicate processing'],
        ['G3 — Coordinate transfers','Search wards for surplus donor, enforce controlled substance policy, execute inter-ward transfer, record audit trail'],
        ['G4 — Initiate procurement','Select approved supplier by drug category, create procurement record, track status lifecycle'],
        ['G5 — Escalate overdue orders','Proactively check elapsed steps every cycle, mark ESCALATED after 5 steps without confirmation'],
    ]
)
h3('Severity Thresholds')
add_table(
    ['Severity', 'Condition (ratio = stock / threshold)', 'Response'],
    [
        ['CRITICAL', 'ratio < 10%',  'Parallel transfer + procurement (with await flag)'],
        ['HIGH',     '10% – 30%',    'Transfer-first; procurement only if transfer fails'],
        ['MEDIUM',   '30% – 50%',    'Direct procurement (no transfer attempted)'],
        ['LOW / OK', '> 50%',        'No action'],
    ]
)
h3('19 Functionalities (F1 – F19)')
funcs = [
    ('F1',  'Receive drug stock readings from the pharmacy sensor network per simulation step'),
    ('F2',  'Compute stock level percentage relative to reorder threshold'),
    ('F3',  'Classify stock level by severity: CRITICAL, HIGH, MEDIUM'),
    ('F4',  'Suppress duplicate alerts for stock-ward pairs already alerted'),
    ('F5',  'Classify shortage by drug category and select response plan'),
    ('F6',  'Route CRITICAL shortages to both transfer and procurement agents'),
    ('F7',  'Route HIGH shortages to transfer coordination only'),
    ('F8',  'Route MEDIUM shortages to procurement only'),
    ('F9',  'Enforce controlled substance policy — no internal transfer of morphine'),
    ('F10', 'Search all wards for a surplus donor ward meeting criteria'),
    ('F11', 'Execute inter-ward stock transfer and update both ward records'),
    ('F12', 'Record every transfer attempt with full audit trail'),
    ('F13', 'Select supplier by drug category from the supplier registry'),
    ('F14', 'Create procurement record and notify procurement pipeline'),
    ('F15', 'Proactively monitor pending procurements every cycle for timeout'),
    ('F16', 'Escalate overdue procurement to ESCALATED status after five steps'),
    ('F17', 'Receive and process supplier confirmation events'),
    ('F18', 'Track resolved shortages to avoid reprocessing'),
    ('F19', 'Provide full audit trail via the simulation activity log'),
]
add_table(['ID', 'Functionality'], funcs)

# ── Phase 2 ───────────────────────────────────────────────────────────────────
h2('Phase 2 — Architectural Design')
body(
    'Phase 2 identifies agent types, their responsibilities, beliefs, and the acquaintance '
    'diagram showing inter-agent message flow.'
)
h3('Agent Types')
add_table(
    ['Agent', 'Core Responsibility', 'Goals'],
    [
        ['StockMonitorAgent',        'Monitor sensor readings; classify severity; suppress duplicates', 'G1.1 – G1.4'],
        ['SupplyAssessmentAgent',    'Classify shortage by category; select and route to response plan', 'G2.1 – G2.4'],
        ['TransferCoordinationAgent','Search for donor ward; enforce controlled substance policy; execute transfer', 'G3.1 – G3.5'],
        ['ProcurementEscalationAgent','Initiate procurement; track orders; escalate after 5-step timeout', 'G4.1 – G4.4, G5.1 – G5.3'],
        ['ExpiryMonitorAgent',       'Scan batch expiry every cycle; enforce FEFO; alert on expired batches', 'Implicit — safety'],
    ]
)
h3('Acquaintance (Message Flow)')
body(
    'PharmacySensor -> StockMonitorAgent (STOCK_READING)\n'
    'StockMonitorAgent -> SupplyAssessmentAgent (STOCK_ALERT)\n'
    'SupplyAssessmentAgent -> TransferCoordinationAgent (SHORTAGE_CLASSIFIED)\n'
    'SupplyAssessmentAgent -> ProcurementEscalationAgent (SHORTAGE_CLASSIFIED, with await flag on CRITICAL)\n'
    'TransferCoordinationAgent -> ProcurementEscalationAgent (TRANSFER_RESULT)\n'
    'ExpiryMonitorAgent -> ProcurementEscalationAgent (EXPIRY_ALERT)'
)
body('Communication is strictly directional. No agent sends messages back to a higher-level agent except via TRANSFER_RESULT.')

h3('Data Stores')
add_table(
    ['Store', 'Managed by', 'Contents'],
    [
        ['DrugDatabase',         'StockMonitorAgent, TransferCoordinationAgent', 'Drug metadata, all ward stock records, expiry batches'],
        ['WardDatabase',         'TransferCoordinationAgent', 'Ward list, all TransferRecords (TR-001 format)'],
        ['SupplierDatabase',     'ProcurementEscalationAgent', 'Supplier registry, all ProcurementRecords (PR-001 format)'],
        ['active_shortages',     'SupplyAssessmentAgent', 'Dict of shortage_id to shortage info (duplicate prevention)'],
        ['pending_procurements', 'ProcurementEscalationAgent', 'Dict tracking dispatch_step and escalation state'],
        ['resolved_shortages',   'ProcurementEscalationAgent', 'Dict of shortage_id to resolution details'],
    ]
)

# ── Phase 3 ───────────────────────────────────────────────────────────────────
h2('Phase 3 — Interaction Design')
body(
    'Phase 3 defines message performatives and interaction diagrams for all key scenarios.'
)
h3('Message Performatives')
add_table(
    ['Performative', 'Direction', 'Key Content Fields'],
    [
        ['STOCK_READING',        'PharmacySensor → StockMonitorAgent', 'drug_id, ward_id, quantity, step'],
        ['STOCK_ALERT',          'StockMonitor → SupplyAssessment', 'drug_id, ward_id, severity, quantity_needed, step'],
        ['SHORTAGE_CLASSIFIED',  'SupplyAssessment → Transfer or Procurement', 'shortage_id, severity, category, await_transfer_result (flag)'],
        ['TRANSFER_RESULT',      'TransferCoord → ProcurementEscalation', 'success (bool), reason, transfer_id, transfer_qty, from_ward'],
        ['EXPIRY_ALERT',         'ExpiryMonitor → ProcurementEscalation', 'drug_id, ward_id, batch_id, expiry_step, status'],
        ['ESCALATION_NOTICE',    'Internal record', 'procurement_id, escalated_step'],
    ]
)

h3('Scenario 1: Critical Insulin Shortage — ICU (Transfer Success)')
add_table(
    ['Step', 'Event'],
    [
        ['Step 1', 'Sensor: Insulin / ICU = 8 units (threshold 100 → 8% → CRITICAL)'],
        ['Step 1', 'StockMonitorAgent sends STOCK_ALERT to SupplyAssessmentAgent'],
        ['Step 1', 'SupplyAssessmentAgent routes SHORTAGE_CLASSIFIED to both Transfer and Procurement agents (await_transfer_result=True to Procurement)'],
        ['Step 1', 'TransferCoordinationAgent finds Emergency ward: 120 units, surplus ratio 1.2 > 0.5; transfers 60 units'],
        ['Step 1', 'TRANSFER_RESULT(success=True, TR-001) sent to ProcurementEscalationAgent'],
        ['Step 1', 'ProcurementEscalationAgent marks shortage RESOLVED — no procurement initiated'],
    ]
)

h3('Scenario 2: Morphine — Controlled Substance Escalation')
add_table(
    ['Step', 'Event'],
    [
        ['Step 3',   'Sensor: Morphine / Surgical = 7 mg (threshold 50 → 14% → HIGH)'],
        ['Step 3',   'SupplyAssessmentAgent routes to TransferCoordinationAgent only (HIGH plan)'],
        ['Step 3',   'TransferCoordinationAgent: CONTROLLED drug — regulatory denial; TRANSFER_RESULT(success=False, TR-002)'],
        ['Step 3',   'ProcurementEscalationAgent creates PR-001 via PharmaControl Ltd; dispatch_step=3'],
        ['Steps 4-7','Proactive check: 8-3 < 5 — no escalation yet'],
        ['Step 8',   'Proactive check: 8-3 = 5 >= timeout — PR-001 status = ESCALATED'],
    ]
)

h3('Scenario 3: Multi-Drug Shortages')
add_table(
    ['Step', 'Drug / Ward', 'Outcome'],
    [
        ['Step 8',  'Amoxicillin / Surgical (40 tabs, 20% → HIGH)',   'ICU donor (150 tabs, ratio 0.75); TR-003 COMPLETED — 75 tabs transferred'],
        ['Step 10', 'Paracetamol / General (160 tabs, 32% → MEDIUM)', 'PR-002 created via BasicMeds (direct procurement, no transfer)'],
        ['Step 14', 'Paracetamol / General',                          'Supplier confirmation arrives — PR-002 status = CONFIRMED'],
    ]
)

# ── Phase 4 ───────────────────────────────────────────────────────────────────
h2('Phase 4 — Detailed Design')
body(
    'Phase 4 specifies capabilities, plans with context conditions, belief structures, '
    'and percept/action tables for each agent.'
)

h3('Plan Selection Logic — SupplyAssessmentAgent')
add_table(
    ['Plan', 'Context Condition', 'Actions'],
    [
        ['CRITICAL_plan', 'severity == CRITICAL', 'Send SHORTAGE_CLASSIFIED to Transfer AND Procurement (await_transfer_result=True)'],
        ['HIGH_plan',     'severity == HIGH',     'Send SHORTAGE_CLASSIFIED to Transfer only'],
        ['MEDIUM_plan',   'severity == MEDIUM',   'Send SHORTAGE_CLASSIFIED to Procurement only (no transfer)'],
    ]
)

h3('Plan Selection Logic — TransferCoordinationAgent')
add_table(
    ['Plan', 'Context Condition', 'Actions'],
    [
        ['ControlledDenialPlan', 'drug.category == CONTROLLED', 'Log denial; create FAILED TransferRecord; send TRANSFER_RESULT(success=False)'],
        ['TransferSearchPlan',   'drug.category != CONTROLLED',
         'Iterate wards; find donor where surplus_ratio > 0.5 AND stock > qty*0.5; '
         'transfer_qty = min(needed, int(donor*0.5)); update both stocks; '
         'if no donor → TRANSFER_RESULT(success=False)'],
    ]
)

h3('Plan Selection Logic — ProcurementEscalationAgent')
add_table(
    ['Trigger / Percept', 'Plan', 'Actions'],
    [
        ['SHORTAGE_CLASSIFIED (await=True)',  'WaitForTransfer',        'Store in _awaiting_transfer dict and hold'],
        ['SHORTAGE_CLASSIFIED (await=False)', 'InitiateProcurement',    'Select supplier by category; create ProcurementRecord; store in pending_procurements'],
        ['TRANSFER_RESULT (success=True)',    'ResolveShortage',        'Move shortage to resolved_shortages; delete from _awaiting_transfer'],
        ['TRANSFER_RESULT (success=False)',   'InitiateProcurement',    'Pop from _awaiting_transfer; call _initiate_procurement()'],
        ['proactive_step() each cycle',       'CheckAndEscalate',       'For each REQUESTED procurement: if current_step - dispatch_step >= 5 → mark ESCALATED'],
        ['simulate_supplier_confirmation()',  'ConfirmSupply',          'Find procurement in pending_procurements; set status CONFIRMED'],
    ]
)

h3('Belief Structures')
add_table(
    ['Module', 'Key Structures'],
    [
        ['drug_belief.py',     'DrugCategory (CONTROLLED/ESSENTIAL/STANDARD), SeverityLevel, Drug, StockRecord.severity(), ExpiryBatch, DrugDatabase'],
        ['ward_belief.py',     'TransferStatus (PENDING/COMPLETED/FAILED), Ward, TransferRecord (TR-001 format), WardDatabase'],
        ['supplier_belief.py', 'ProcurementStatus (REQUESTED/CONFIRMED/ESCALATED), Supplier, ProcurementRecord (PR-001 format), SupplierDatabase'],
    ]
)

# ── Phase 5 ───────────────────────────────────────────────────────────────────
h2('Phase 5 — Implementation Report')
body(
    'Phase 5 maps every design decision from Phases 1–4 to the Python implementation '
    'and evaluates challenges and limitations.'
)
h3('Platform and Implementation')
bullet('Python 3, standard library only (dataclasses, enums, collections.deque, tkinter)')
bullet('No third-party packages required — runs on any university laboratory computer')
bullet('Five agents in src/agents/, three belief modules in src/beliefs/')
bullet('Base Agent class: inbox deque, send/receive interface, run_cycle(), proactive_step()')
bullet('AgentSystem registry: name-based message routing between agents')
bullet('MedStockSimulator: 20-step loop, pre-scheduled sensor events, UI callbacks')

h3('Scenario Verification')
add_table(
    ['Scenario', 'Expected', 'Actual'],
    [
        ['1 — Insulin ICU',          'TR-001 COMPLETED, no procurement',         'Verified at step 1'],
        ['2 — Morphine Surgical',    'TR-002 FAILED (controlled), PR-001 ESCALATED at step 8', 'Verified at steps 3–8'],
        ['3a — Amoxicillin Surgical','TR-003 COMPLETED, no procurement',         'Verified at step 8'],
        ['3b — Paracetamol General', 'PR-002 CONFIRMED at step 14',              'Verified at step 14'],
    ]
)

h3('Challenges and Limitations')
body(
    'The most significant design challenge was handling CRITICAL shortages where both agents '
    'receive SHORTAGE_CLASSIFIED simultaneously. The solution uses the _awaiting_transfer '
    'dictionary: the ProcurementEscalationAgent stores the content and waits; it only calls '
    '_initiate_procurement() when TRANSFER_RESULT arrives with success=False.'
)
body(
    'Primary limitation: the system operates in a closed simulation. PharmacySensor fires '
    'pre-scheduled events; supplier confirmation is injected directly by the simulator. '
    'Extending to a live hospital information system would require replacing PharmacySensor '
    'with a network-connected adapter and adding a webhook listener — both architecturally '
    'straightforward given the agent separation of concerns.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — AGENT PROFILES
# ══════════════════════════════════════════════════════════════════════════════
h1('3. Agent Profiles')
divider()

agents = [
    {
        'name': '3.1 StockMonitorAgent',
        'role': 'Entry point for all environmental data. Detects stock shortages.',
        'goals': 'G1.1 – G1.4',
        'receives': 'STOCK_READING from PharmacySensor',
        'sends': 'STOCK_ALERT to SupplyAssessmentAgent',
        'beliefs': 'DrugDatabase (stock levels, thresholds, drug metadata); alerted_stocks set',
        'key_logic': [
            'ratio = current_stock / reorder_threshold',
            'CRITICAL: ratio < 0.10 | HIGH: 0.10–0.30 | MEDIUM: 0.30–0.50',
            'alerted_stocks set prevents duplicate alerts for same (drug, ward) pair',
        ],
    },
    {
        'name': '3.2 SupplyAssessmentAgent',
        'role': 'Central classifier and router. Selects the correct response plan.',
        'goals': 'G2.1 – G2.4',
        'receives': 'STOCK_ALERT from StockMonitorAgent',
        'sends': 'SHORTAGE_CLASSIFIED to TransferCoordinationAgent and/or ProcurementEscalationAgent',
        'beliefs': 'DrugDatabase (drug category); active_shortages dict',
        'key_logic': [
            'CRITICAL severity → send to BOTH agents (await_transfer_result=True to Procurement)',
            'HIGH severity → send to TransferCoordinationAgent only',
            'MEDIUM severity → send to ProcurementEscalationAgent only',
            'active_shortages dict prevents re-classifying the same shortage',
        ],
    },
    {
        'name': '3.3 TransferCoordinationAgent',
        'role': 'Executes inter-ward stock transfers. Enforces controlled substance policy.',
        'goals': 'G3.1 – G3.5',
        'receives': 'SHORTAGE_CLASSIFIED from SupplyAssessmentAgent',
        'sends': 'TRANSFER_RESULT to ProcurementEscalationAgent',
        'beliefs': 'DrugDatabase (all ward stocks); WardDatabase (ward list, transfer records)',
        'key_logic': [
            'CONTROLLED drug → immediate denial (regulatory policy), no search performed',
            'For non-controlled: iterate all wards except shortage ward',
            'Donor valid if surplus_ratio > 0.5 AND stock > quantity_needed * 0.5',
            'transfer_qty = min(quantity_needed, int(donor_stock * 0.5))',
            'Updates both donor and recipient stocks in DrugDatabase',
        ],
    },
    {
        'name': '3.4 ProcurementEscalationAgent',
        'role': 'Manages full procurement lifecycle. Proactively escalates overdue orders.',
        'goals': 'G4.1 – G4.4, G5.1 – G5.3',
        'receives': 'SHORTAGE_CLASSIFIED (from SupplyAssessment), TRANSFER_RESULT (from Transfer), EXPIRY_ALERT (from ExpiryMonitor)',
        'sends': 'Creates ProcurementRecords, updates SupplierDatabase',
        'beliefs': 'SupplierDatabase; pending_procurements dict; resolved_shortages dict; _awaiting_transfer dict',
        'key_logic': [
            'On SHORTAGE_CLASSIFIED with await=True: store and wait for transfer outcome',
            'On TRANSFER_RESULT success=False: call _initiate_procurement()',
            'Every cycle — proactive_step(): check elapsed steps for all REQUESTED procurements',
            'ESCALATION_TIMEOUT_STEPS = 5: if current_step - dispatch_step >= 5 → ESCALATED',
            'Supplier confirmation: simulate_supplier_confirmation() sets status to CONFIRMED',
        ],
    },
    {
        'name': '3.5 ExpiryMonitorAgent',
        'role': 'Monitors batch expiry. Enforces FEFO (First Expired, First Out) ordering.',
        'goals': 'Implicit — medication safety and waste reduction',
        'receives': '(No incoming messages — purely proactive)',
        'sends': 'EXPIRY_ALERT to ProcurementEscalationAgent',
        'beliefs': 'DrugDatabase (all expiry batches across all wards)',
        'key_logic': [
            'Runs proactively every cycle — no message trigger required',
            'EXPIRING_SOON_THRESHOLD = 3 steps; WARNING_THRESHOLD = 6 steps',
            'steps_remaining = expiry_step - current_step',
            'FEFO: _consume_batches_fefo() consumes earliest-expiring batches first',
            '_register_restock_batch() creates new batch when stock increases',
        ],
    },
]

for ag in agents:
    h2(ag['name'])
    labelled('Role', ag['role'])
    labelled('Goals', ag['goals'])
    labelled('Receives', ag['receives'])
    labelled('Sends', ag['sends'])
    labelled('Beliefs', ag['beliefs'])
    h3('Key Logic')
    for kl in ag['key_logic']:
        bullet(kl)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — SYSTEM ARCHITECTURE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1('4. System Architecture Summary')
divider()

h2('Code Structure')
add_table(
    ['Module', 'Path', 'Contents'],
    [
        ['Core framework',   'src/core/',        'message.py, agent.py, agent_system.py'],
        ['Agents',           'src/agents/',       '5 agent Python files'],
        ['Belief stores',    'src/beliefs/',      'drug_belief.py, ward_belief.py, supplier_belief.py'],
        ['Environment',      'src/environment/',  'pharmacy_sensor.py'],
        ['Simulation',       'src/simulation/',   'simulator.py'],
        ['UI',               'ui/',               'Tkinter app.py'],
        ['Documentation',    'docs/',             'phase1 – phase5 markdown files, presentation guide'],
        ['Diagrams',         'diagrams/',         '10 ASCII art diagram files'],
    ]
)

h2('Key Constants')
add_table(
    ['Constant', 'Value', 'Effect'],
    [
        ['ESCALATION_TIMEOUT_STEPS', '5',  'Steps without confirmation before procurement is escalated'],
        ['EXPIRING_SOON_THRESHOLD',  '3',  'Steps remaining before a batch is flagged EXPIRING_SOON'],
        ['WARNING_THRESHOLD',        '6',  'Steps remaining before a batch is flagged WARNING'],
        ['TOTAL_STEPS',              '20', 'Duration of each simulation run'],
        ['Surplus ratio minimum',    '0.5','Donor ward must have surplus_ratio > 0.5 to be eligible'],
    ]
)

h2('Conclusion')
body(
    'MedStock demonstrates that the Prometheus methodology produces a coherent, verifiable, '
    'and extensible multi-agent system. The four-phase design process from system specification '
    'through detailed design produced a direct and auditable mapping to the Python implementation '
    'with no design decisions left unaccounted for. The system correctly handles all three '
    'simulation scenarios, enforces the controlled substance regulatory policy, performs proactive '
    'escalation monitoring every cycle, tracks batch expiry with FEFO, and provides a professional '
    'graphical interface for visualising and interacting with the simulation in real time.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out = r'c:\Users\hp\OneDrive\Desktop\projects\DCIT403_Project_11126585\MedStock_Report.docx'
doc.save(out)
print(f'Saved: {out}')
